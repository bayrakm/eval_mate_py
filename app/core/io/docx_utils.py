"""
DOCX Visual Extraction Utilities

This module provides DOCX-specific utilities for extracting images, tables, and
equations from Word documents using python-docx.

Supports comprehensive visual content extraction from DOCX files.
"""

import logging
from typing import Iterator, Optional, List, Tuple
from dataclasses import dataclass
from pathlib import Path
from PIL import Image
import io
import re

logger = logging.getLogger(__name__)

# Optional imports with graceful fallbacks
try:
    from docx import Document
    from docx.shared import Inches
    from docx.enum.table import WD_TABLE_ALIGNMENT
    PYTHON_DOCX_AVAILABLE = True
except ImportError:
    PYTHON_DOCX_AVAILABLE = False
    logger.error("python-docx not available, DOCX processing will be disabled")


@dataclass
class DocxImage:
    """
    Represents an extracted image from a DOCX file with metadata.
    
    Attributes:
        filename: Original filename or generated name for the image
        pil_image: PIL Image object of the extracted image
        alt_text: Alternative text if available
        paragraph_index: Index of paragraph containing the image (for context)
    """
    filename: str
    pil_image: Image.Image
    alt_text: Optional[str] = None
    paragraph_index: Optional[int] = None


def iter_docx_images(path: str) -> Iterator[DocxImage]:
    """
    Yield images from DOCX file with any available alt text.
    
    Supports both inline and floating shapes.
    
    Args:
        path: Path to DOCX file
        
    Yields:
        DocxImage objects with PIL images and metadata
    """
    if not PYTHON_DOCX_AVAILABLE:
        logger.error("python-docx not available, cannot extract images")
        return
    
    try:
        doc = Document(path)
        
        # Track paragraph index for context
        for para_idx, paragraph in enumerate(doc.paragraphs):
            # Check for inline images in runs
            for run in paragraph.runs:
                if run.element.xml:
                    # Look for drawing elements that contain images
                    if 'w:drawing' in run.element.xml or 'pic:pic' in run.element.xml:
                        images_in_run = extract_images_from_run(run, para_idx)
                        for img in images_in_run:
                            yield img
        
        # Also check document parts for embedded images
        for rel in doc.part.rels.values():
            if "image" in rel.target_ref:
                try:
                    image_part = rel.target_part
                    image_data = image_part.blob
                    
                    # Create PIL image
                    pil_image = Image.open(io.BytesIO(image_data))
                    
                    # Generate filename from relationship
                    filename = f"embedded_{rel.rId}.png"
                    
                    docx_image = DocxImage(
                        filename=filename,
                        pil_image=pil_image,
                        alt_text=None,  # Alt text not easily accessible from relations
                        paragraph_index=None
                    )
                    yield docx_image
                    
                except Exception as e:
                    logger.warning(f"Failed to extract embedded image {rel.rId}: {e}")
                    continue
                    
    except Exception as e:
        logger.error(f"Failed to extract images from DOCX {path}: {e}")


def extract_images_from_run(run, paragraph_index: int) -> List[DocxImage]:
    """
    Extract images from a specific run element.
    
    Args:
        run: docx Run object
        paragraph_index: Index of containing paragraph
        
    Returns:
        List of DocxImage objects
    """
    images = []
    
    try:
        # Look for drawing elements in the run
        drawings = run.element.xpath('.//w:drawing')
        
        for drawing in drawings:
            # Find inline or anchor elements
            inlines = drawing.xpath('.//wp:inline')
            anchors = drawing.xpath('.//wp:anchor')
            
            all_shapes = inlines + anchors
            
            for shape in all_shapes:
                try:
                    # Look for picture elements
                    pics = shape.xpath('.//pic:pic', namespaces={
                        'pic': 'http://schemas.openxmlformats.org/drawingml/2006/picture'
                    })
                    
                    for pic in pics:
                        # Try to extract image data
                        blips = pic.xpath('.//a:blip', namespaces={
                            'a': 'http://schemas.openxmlformats.org/drawingml/2006/main'
                        })
                        
                        for blip in blips:
                            embed_id = blip.get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed')
                            
                            if embed_id:
                                try:
                                    # Get the image part
                                    image_part = run.part.related_parts[embed_id]
                                    image_data = image_part.blob
                                    
                                    # Create PIL image
                                    pil_image = Image.open(io.BytesIO(image_data))
                                    
                                    # Try to get alt text
                                    alt_text = extract_alt_text_from_pic(pic)
                                    
                                    # Generate filename
                                    filename = f"para_{paragraph_index}_{embed_id}.png"
                                    
                                    docx_image = DocxImage(
                                        filename=filename,
                                        pil_image=pil_image,
                                        alt_text=alt_text,
                                        paragraph_index=paragraph_index
                                    )
                                    images.append(docx_image)
                                    
                                except Exception as e:
                                    logger.warning(f"Failed to extract image with embed_id {embed_id}: {e}")
                                    continue
                                    
                except Exception as e:
                    logger.warning(f"Failed to process shape in paragraph {paragraph_index}: {e}")
                    continue
                    
    except Exception as e:
        logger.warning(f"Failed to extract images from run in paragraph {paragraph_index}: {e}")
    
    return images


def extract_alt_text_from_pic(pic_element) -> Optional[str]:
    """
    Extract alternative text from a picture element.
    
    Args:
        pic_element: XML element representing the picture
        
    Returns:
        Alt text if found, None otherwise
    """
    try:
        # Look for description in various possible locations
        desc_elements = pic_element.xpath('.//pic:cNvPr', namespaces={
            'pic': 'http://schemas.openxmlformats.org/drawingml/2006/picture'
        })
        
        for desc in desc_elements:
            descr = desc.get('descr')
            if descr:
                return descr
                
            title = desc.get('title')
            if title:
                return title
                
        # Alternative location for alt text
        alt_elements = pic_element.xpath('.//a:extLst//a:ext//a14:imgProps//a14:imgLayer//a14:imgEffect', namespaces={
            'a': 'http://schemas.openxmlformats.org/drawingml/2006/main',
            'a14': 'http://schemas.microsoft.com/office/drawing/2010/main'
        })
        
        for alt_elem in alt_elements:
            alt_text = alt_elem.get('alt')
            if alt_text:
                return alt_text
                
    except Exception as e:
        logger.debug(f"Failed to extract alt text: {e}")
    
    return None


def iter_docx_tables(path: str) -> Iterator[List[List[str]]]:
    """
    Yield tables from DOCX as list-of-rows with cells as plain text.
    
    Args:
        path: Path to DOCX file
        
    Yields:
        Tables as list of rows, where each row is a list of cell strings
    """
    if not PYTHON_DOCX_AVAILABLE:
        logger.error("python-docx not available, cannot extract tables")
        return
    
    try:
        doc = Document(path)
        
        for table in doc.tables:
            table_data = []
            
            for row in table.rows:
                row_data = []
                
                for cell in row.cells:
                    # Extract text from all paragraphs in the cell
                    cell_text = ""
                    for paragraph in cell.paragraphs:
                        if cell_text:
                            cell_text += "\n"
                        cell_text += paragraph.text.strip()
                    
                    row_data.append(cell_text)
                
                table_data.append(row_data)
            
            if table_data:  # Only yield non-empty tables
                yield table_data
                
    except Exception as e:
        logger.error(f"Failed to extract tables from DOCX {path}: {e}")


def get_docx_equation_runs(path: str) -> Iterator[str]:
    """
    Best-effort detection of Office Math (OMML) equations in DOCX.
    
    python-docx doesn't expose OMML directly, so we inspect the underlying XML
    for math elements and extract adjacent text as a placeholder.
    
    Args:
        path: Path to DOCX file
        
    Yields:
        Text content that appears to be mathematical equations
    """
    if not PYTHON_DOCX_AVAILABLE:
        logger.error("python-docx not available, cannot extract equations")
        return
    
    try:
        doc = Document(path)
        
        for paragraph in doc.paragraphs:
            # Check if paragraph contains math elements
            para_xml = paragraph._element.xml
            
            # Look for Office Math markup
            if 'm:oMath' in para_xml or 'm:oMathPara' in para_xml:
                # Extract the text content - this is a simplified approach
                text = paragraph.text.strip()
                if text:
                    yield text
                continue
            
            # Also check individual runs for math content
            for run in paragraph.runs:
                run_xml = run.element.xml
                
                if 'm:oMath' in run_xml or 'm:r' in run_xml:
                    # This run contains math - extract text
                    text = run.text.strip()
                    if text:
                        yield text
                        
    except Exception as e:
        logger.warning(f"Failed to extract equations from DOCX {path}: {e}")


def extract_docx_paragraphs_with_context(path: str) -> List[dict]:
    """
    Extract paragraphs from DOCX with context information for caption detection.
    
    Args:
        path: Path to DOCX file
        
    Returns:
        List of dicts with 'text', 'index', 'style', and 'has_image' keys
    """
    paragraphs = []
    
    if not PYTHON_DOCX_AVAILABLE:
        logger.error("python-docx not available, cannot extract paragraphs")
        return paragraphs
    
    try:
        doc = Document(path)
        
        for idx, paragraph in enumerate(doc.paragraphs):
            # Check if paragraph contains images
            has_image = False
            para_xml = paragraph._element.xml
            if 'w:drawing' in para_xml or 'pic:pic' in para_xml:
                has_image = True
            
            # Get paragraph style
            style_name = paragraph.style.name if paragraph.style else "Normal"
            
            paragraph_info = {
                'text': paragraph.text.strip(),
                'index': idx,
                'style': style_name,
                'has_image': has_image
            }
            
            paragraphs.append(paragraph_info)
            
    except Exception as e:
        logger.warning(f"Failed to extract paragraphs from DOCX {path}: {e}")
    
    return paragraphs


def detect_figure_captions_in_docx(paragraphs: List[dict]) -> List[Tuple[int, str]]:
    """
    Detect figure captions in DOCX paragraphs using heuristics.
    
    Args:
        paragraphs: List of paragraph info dicts from extract_docx_paragraphs_with_context
        
    Returns:
        List of tuples (paragraph_index, caption_text)
    """
    captions = []
    
    # Caption patterns
    caption_patterns = [
        r'^Figure\s+\d+[:\.\-]\s*(.+)',
        r'^Fig\.\s*\d+[:\.\-]\s*(.+)',
        r'^Table\s+\d+[:\.\-]\s*(.+)',
        r'^Image\s+\d+[:\.\-]\s*(.+)',
        r'^Diagram\s+\d+[:\.\-]\s*(.+)',
    ]
    
    for para in paragraphs:
        text = para['text']
        index = para['index']
        
        if not text:
            continue
        
        # Check if this looks like a caption
        for pattern in caption_patterns:
            match = re.match(pattern, text, re.IGNORECASE)
            if match:
                caption_text = match.group(1).strip() if match.groups() else text
                captions.append((index, caption_text))
                break
    
    return captions


def is_equation_paragraph(paragraph_text: str) -> bool:
    """
    Heuristic to determine if a paragraph contains mathematical content.
    
    Args:
        paragraph_text: Text content of the paragraph
        
    Returns:
        True if paragraph appears to contain equations
    """
    # Import OCR module for math detection
    try:
        from app.core.io.ocr import detect_math_symbols
        return detect_math_symbols(paragraph_text)
    except ImportError:
        # Fallback to simple pattern matching
        math_indicators = ['=', '+', '-', '×', '÷', '∑', '∫', '√', '²', '³', 'x²', 'x³']
        text_lower = paragraph_text.lower()
        
        # Check for multiple math indicators
        math_count = sum(1 for indicator in math_indicators if indicator in text_lower)
        return math_count >= 2